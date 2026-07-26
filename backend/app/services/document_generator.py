# backend/app/services/document_generator.py
from docx import Document
import tempfile
import os
import re
from typing import Dict, Any, List, Tuple, Optional
import subprocess
import platform
from datetime import datetime 
import logging

# Configure logging
logger = logging.getLogger(__name__)

class DocumentGenerator:
    def __init__(self):
        self.placeholder_pattern = r'\{\{(\w+)\}\}'
        
        # Check PDF conversion capability
        self.pdf_available = self._check_pdf_conversion()
    
    def _check_pdf_conversion(self) -> bool:
        """Check if PDF conversion is available"""
        try:
            if platform.system() == "Windows":
                # Try docx2pdf (requires Microsoft Word)
                import docx2pdf
                return True
            else:
                # Try LibreOffice
                result = subprocess.run(['libreoffice', '--version'], 
                    capture_output=True, text=True, timeout=5)
                return result.returncode == 0
        except Exception as e:
            logger.error(f"PDF conversion check failed: {e}")
            return False
    
    def generate_document(self, template_path: str, user_data: Dict[str, Any], 
            output_format: str = "docx") -> Tuple[str, Dict[str, Any]]:
        """Generate a document from template and user data
        
        Returns:
            Tuple of (file_path, metadata)
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
            
        try:
            # Load the template document
            doc = Document(template_path)
            
            # Process all content with improved placeholder handling
            replacements_made = 0
            replacements_made += self._process_paragraphs(doc.paragraphs, user_data)
            replacements_made += self._process_tables(doc.tables, user_data)
            replacements_made += self._process_headers_footers(doc.sections, user_data)
            
            # Generate base filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            base_filename = f"generated_{timestamp}"
            
            # Save as DOCX first
            temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_docx.name)
            temp_docx.close()
            
            # Convert to PDF if requested
            if output_format.lower() == "pdf":
                if not self.pdf_available:
                    raise Exception("PDF conversion not available on this system")
                
                temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                temp_pdf.close()
                
                success = self._convert_to_pdf(temp_docx.name, temp_pdf.name)
                if not success:
                    raise Exception("PDF conversion failed")
                
                # Clean up DOCX and return PDF
                os.unlink(temp_docx.name)
                final_file = temp_pdf.name
            else:
                final_file = temp_docx.name
            
            # Get file metadata
            metadata = {
                "replacements_made": replacements_made,
                "file_size": os.path.getsize(final_file),
                "format": output_format,
                "mime_type": self._get_mime_type(output_format)
            }
            
            return final_file, metadata
            
        except Exception as e:
            # Cleanup any temporary files
            for temp_file in ["temp_docx", "temp_pdf"]:
                path = locals().get(temp_file)
                if path:
                    if hasattr(path, "name"):
                        file_path = path.name
                    else:
                        file_path = path
                    if os.path.exists(file_path):
                        try:
                            os.unlink(file_path)
                        except Exception as cleanup_err:
                            logger.error(f"Failed to clean up {file_path}: {cleanup_err}")
            raise Exception(f"Document generation failed: {str(e)}")
    
    def _convert_to_pdf(self, docx_path: str, pdf_path: str) -> bool:
        """Convert DOCX to PDF"""
        try:
            if platform.system() == "Windows":
                # Use docx2pdf on Windows
                from docx2pdf import convert
                convert(docx_path, pdf_path)
                return True
            else:
                # Use LibreOffice on Linux/Mac
                output_dir = os.path.dirname(pdf_path)
                result = subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', output_dir, docx_path
                ], capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0:
                    # LibreOffice creates PDF with same basename
                    generated_pdf = os.path.join(
                        output_dir, 
                        os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
                    )
                    if os.path.exists(generated_pdf):
                        # Move to desired location
                        if generated_pdf != pdf_path:
                            os.rename(generated_pdf, pdf_path)
                        return True
                return False
        except Exception as e:
            logger.error(f"PDF conversion error: {e}")
            return False
    
    def _get_mime_type(self, format_type: str) -> str:
        """Get MIME type for format"""
        mime_types = {
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "pdf": "application/pdf"
        }
        return mime_types.get(format_type.lower(), "application/octet-stream")
    
    def _process_paragraphs(self, paragraphs, user_data: Dict[str, Any]) -> int:
        """Process paragraphs and replace placeholders while preserving formatting"""
        replacements = 0
        for paragraph in paragraphs:
            # Get full paragraph text to check for placeholders
            full_text = paragraph.text
            if not self._contains_placeholders(full_text):
                continue
            
            # Use the method to handle split placeholders
            replacements += self._replace_split_placeholders(paragraph, user_data)
        return replacements
    
    def _replace_split_placeholders(self, paragraph, user_data: Dict[str, Any]) -> int:
        """Replace placeholders that might be split across multiple runs"""
        full_text = paragraph.text
        
        # Find all placeholders in the complete text
        placeholders = list(re.finditer(self.placeholder_pattern, full_text))
        if not placeholders:
            return 0
        
        # Process placeholders from right to left to maintain positions
        for match in reversed(placeholders):
            placeholder_name = match.group(1)
            replacement_value = str(user_data.get(placeholder_name, f"{{{{{placeholder_name}}}}}"))
            start_pos, end_pos = match.span()
            
            # Replace the text across runs
            self._replace_text_across_runs(paragraph, start_pos, end_pos, replacement_value)
        
        return len(placeholders)
    
    def _replace_text_across_runs(self, paragraph, start_pos: int, end_pos: int, replacement_text: str):
        """Replace text that spans across multiple runs"""
        # Build a map of character positions to runs
        runs_info = []
        current_pos = 0
        
        for run_index, run in enumerate(paragraph.runs):
            run_length = len(run.text)
            runs_info.append({
                'run': run,
                'run_index': run_index,
                'start_pos': current_pos,
                'end_pos': current_pos + run_length,
                'original_text': run.text
            })
            current_pos += run_length
        
        # Find which runs are affected by the replacement
        affected_runs = []
        for run_info in runs_info:
            # Check if this run overlaps with the replacement range
            if run_info['start_pos'] < end_pos and run_info['end_pos'] > start_pos:
                affected_runs.append(run_info)
        
        if not affected_runs:
            return
        
        # Calculate what part of each run to keep/replace
        for i, run_info in enumerate(affected_runs):
            run = run_info['run']
            run_start = run_info['start_pos']
            run_end = run_info['end_pos']
            original_text = run_info['original_text']
            
            # Calculate the slice positions within this run
            slice_start = max(0, start_pos - run_start)
            slice_end = min(len(original_text), end_pos - run_start)
            
            if i == 0:
                # First affected run: keep text before placeholder + replacement
                new_text = original_text[:slice_start] + replacement_text + original_text[slice_end:]
                run.text = new_text
            else:
                # Other affected runs: keep text after placeholder (if any)
                if slice_end < len(original_text):
                    run.text = original_text[slice_end:]
                else:
                    run.text = ""
    
    def _process_tables(self, tables, user_data: Dict[str, Any]) -> int:
        """Process tables and replace placeholders"""
        replacements = 0
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replacements += self._process_paragraphs(cell.paragraphs, user_data)
        return replacements
    
    def _process_headers_footers(self, sections, user_data: Dict[str, Any]) -> int:
        """Process headers and footers"""
        replacements = 0
        for section in sections:
            # Process header
            if section.header:
                replacements += self._process_paragraphs(section.header.paragraphs, user_data)
            
            # Process footer  
            if section.footer:
                replacements += self._process_paragraphs(section.footer.paragraphs, user_data)
        return replacements
    
    def _contains_placeholders(self, text: str) -> bool:
        """Check if text contains placeholders"""
        return bool(re.search(self.placeholder_pattern, text))
    
    def generate_both_formats(self, template_path: str, user_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate both DOCX and PDF versions if possible"""
        results = {}
        
        # Generate DOCX
        try:
            docx_path, docx_metadata = self.generate_document(template_path, user_data, "docx")
            results["docx"] = {
                "success": True,
                "file_path": docx_path,
                "metadata": docx_metadata
            }
        except Exception as e:
            results["docx"] = {
                "success": False,
                "error": str(e)
            }
        
        # Generate PDF if available
        if self.pdf_available:
            try:
                pdf_path, pdf_metadata = self.generate_document(template_path, user_data, "pdf")
                results["pdf"] = {
                    "success": True,
                    "file_path": pdf_path,
                    "metadata": pdf_metadata
                }
            except Exception as e:
                results["pdf"] = {
                    "success": False,
                    "error": str(e)
                }
        else:
            results["pdf"] = {
                "success": False,
                "error": "PDF conversion not available"
            }
        
        return results
    
    def preview_replacements(self, template_path: str, user_data: Dict[str, Any]) -> Dict[str, Any]:
        """Preview what replacements will be made (for debugging)"""
        try:
            doc = Document(template_path)
            replacements = []
            
            # Check paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                original = paragraph.text
                if self._contains_placeholders(original):
                    # Simulate replacement
                    def replace_match(match):
                        key = match.group(1)
                        return str(user_data.get(key, f"{{{{ {key} }}}}"))
                    
                    replaced = re.sub(self.placeholder_pattern, replace_match, original)
                    if original != replaced:
                        replacements.append({
                            "type": "paragraph",
                            "index": i,
                            "original": original,
                            "replaced": replaced
                        })
            
            return {
                "success": True,
                "replacements": replacements,
                "total_replacements": len(replacements),
                "pdf_available": self.pdf_available
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "replacements": [],
                "total_replacements": 0,
                "pdf_available": self.pdf_available
            }