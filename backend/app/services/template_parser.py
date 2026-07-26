from docx import Document
import re
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)

class TemplateParser:
    def __init__(self):
        # support camelCase, snake_case, row.table_field
        self.placeholder_pattern = r'\{\{\s*([\w\.]+)\s*\}\}'
        
        self.field_groups = {
            "header": ["nomor", "number", "tanggal", "date", "lampiran", "attachment", "hal", "subject", "perihal", "bidang"],
            "recipient": ["kepada", "recipient", "yth", "alamat", "address", "kota", "location", "tempat"],
            "personal": ["nama", "nim", "id", "program", "student", "mahasiswa", "prodi", "ttl"],
            "parent": ["ortu", "orangtua", "ayah", "ibu", "wali"],
            "content": ["judul", "title", "kegiatan", "activity", "penelitian", "research", "lama", "duration", "waktu", "period", "lokasi", "isi", "tanggal"],
            "signature": ["penandatangan", "signer", "nip", "jabatan", "position", "direktur", "kepala"],
            "other": []
        }

    def parse_template(self, file_path: str) -> Dict[str, Any]:
        try:
            doc = Document(file_path)
            placeholders = self._extract_placeholders(doc)
            schema = self._generate_schema(placeholders)

            return {
                "success": True,
                "placeholders": placeholders,
                "schema": schema,
                "field_count": len(placeholders)
            }
        except Exception as e:
            logger.error(f"Failed to parse template {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "placeholders": [],
                "schema": {"sections": []},
                "field_count": 0
            }

    def _extract_placeholders(self, doc: Document) -> List[str]:
        placeholders = set()
        try:
            for paragraph in doc.paragraphs:
                matches = re.findall(self.placeholder_pattern, paragraph.text)
                placeholders.update(matches)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = re.findall(self.placeholder_pattern, paragraph.text)
                            placeholders.update(matches)
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    matches = re.findall(self.placeholder_pattern, paragraph.text)
                    placeholders.update(matches)
                for paragraph in section.footer.paragraphs:
                    matches = re.findall(self.placeholder_pattern, paragraph.text)
                    placeholders.update(matches)
        except Exception as e:
            logger.warning(f"Error extracting placeholders: {e}")
        return sorted(list(placeholders))

    def _humanize_field(self, field_name: str) -> str:
        replacements = {
            "nim": "NIM",
            "nip": "NIP",
            "ttl": "Tempat, Tanggal Lahir",
            "prodi": "Program Studi",
            "nama": "Nama",
            "nomor": "Nomor",
            "hal": "Hal",
            "ortu": "Orang Tua",
            "bidang": "Nomor"
        }

        base_name = field_name.split(".")[-1]
        lower_name = base_name.lower()

        if lower_name in replacements:
            return replacements[lower_name]

        if "_" in base_name:
            return base_name.replace("_", " ").title()
        label = re.sub(r'([a-z])([A-Z])', r'\1 \2', base_name)
        return label.title()

    def _translate_section_name(self, section_name: str) -> str:
        translations = {
            "header": "Kop Surat",
            "recipient": "Penerima",
            "personal": "Data Mahasiswa",
            "parent": "Data Orang Tua",
            "content": "Isi Surat",
            "signature": "Penandatangan",
            "table": "Tabel Data",
            "other": "Lainnya"
        }
        return translations.get(section_name, section_name.title())

    def _infer_field_type(self, field_name: str) -> str:
        name = field_name.lower()
        
        if name.startswith("row."):
            return "text"

        if any(k in name for k in ["tanggal", "date", "tgl", "ttl"]):
            return "date"
        if "email" in name:
            return "email"
        if any(k in name for k in ["nomor", "number", "nim", "nip", "id", "bidang"]):
            return "number"
        if any(k in name for k in ["alamat", "address", "isi", "keterangan", "description", "perihal"]):
            return "textarea"
        if any(k in name for k in ["nama", "title", "jabatan", "lokasi", "program", "prodi", "kegiatan", "penandatangan"]):
            return "text"
        
        return "text"

    def _generate_schema(self, placeholders: list) -> dict:
        schema = {"sections": []}
        sections: Dict[str, List[Dict[str, Any]]] = {}

        for field in placeholders:
            field_info = {
                "name": field,
                "label": self._humanize_field(field),
                "type": self._infer_field_type(field),
                "repeatable": field.startswith("row.")
            }

            if field.startswith("row."):
                sections.setdefault("table", []).append(field_info)
                continue

            added = False
            for section_name, keywords in self.field_groups.items():
                # khusus tanggal → bedakan
                if "tanggal" in field.lower() and field.lower() != "tanggal":
                    section_name = "content"
                if any(k in field.lower() for k in keywords):
                    sections.setdefault(section_name, []).append(field_info)
                    added = True
                    break

            if not added:
                sections.setdefault("other", []).append(field_info)

        for section_name, fields in sections.items():
            schema["sections"].append({
                "name": self._translate_section_name(section_name),
                "fields": fields
            })

        return schema
